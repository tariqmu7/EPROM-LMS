# -*- coding: utf-8 -*-
"""Generate the EPROM Academy proposal / requirements Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x00, 0x52, 0x9B)
LIME = RGBColor(0x6A, 0xA8, 0x1F)
INK = RGBColor(0x0F, 0x24, 0x38)
MUTED = RGBColor(0x5A, 0x6F, 0x84)

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = INK


def shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tc.append(sh)


def h(text, size=16, color=BLUE, before=14, after=6, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def eyebrow(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = LIME
    return p


def body(text, after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet" + (" 2" if sub else ""))
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade(hdr[i], "00529B")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, wdt in enumerate(widths):
                row.cells[i].width = Inches(wdt)
    return t


# ---------------------------------------------------------------- TITLE PAGE
title_bar = doc.add_paragraph()
title_bar.paragraph_format.space_after = Pt(2)
r = title_bar.add_run("EPROM")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = BLUE
r2 = title_bar.add_run("  ACADEMY")
r2.bold = True
r2.font.size = Pt(30)
r2.font.color.rgb = LIME

sub = doc.add_paragraph()
r = sub.add_run("Interactive Training Platform")
r.font.size = Pt(18)
r.font.color.rgb = INK
r.bold = True

doc.add_paragraph()
tag = doc.add_paragraph()
r = tag.add_run("Proposal & Build Requirements")
r.font.size = Pt(14)
r.font.color.rgb = MUTED

doc.add_paragraph()
quote = doc.add_paragraph()
r = quote.add_run("“Committed to Energy. Built on Trust.”")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = BLUE

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run(
    "Replacing PDFs, PowerPoint and physical site visits with one interactive, "
    "self-paced learning experience — featuring 3D equipment, 360° virtual "
    "site visits, live simulations and built-in assessment."
).font.color.rgb = MUTED

doc.add_paragraph()
doc.add_paragraph("Prepared for: EPROM Training Center Management").runs[0].font.color.rgb = MUTED
doc.add_paragraph("Accompanies: working software prototype (this repository)").runs[0].font.color.rgb = MUTED

doc.add_page_break()

# ---------------------------------------------------------------- 1. SUMMARY
eyebrow("Section 1")
h("Executive Summary")
body(
    "EPROM Academy is a web-based, interactive training platform for the oil & gas "
    "domain. It consolidates the materials our training center uses today — PowerPoint "
    "decks, PDF manuals and costly physical site visits — into a single, self-paced "
    "interactive lesson that the trainee actively works through rather than passively reads."
)
body(
    "Inside one lesson a trainee can read brand-styled content, watch embedded video, "
    "explore equipment in interactive 3D, take a 360° virtual walk-down of a plant area, "
    "run a live equipment simulation, and prove their understanding with quizzes and "
    "hands-on tasks — with their progress and scores tracked throughout."
)
body(
    "A working prototype has already been built (delivered alongside this document) that "
    "demonstrates every one of these capabilities end-to-end on a sample course, "
    "“Centrifugal Pump: Operation & Maintenance.” This document explains the concept, "
    "the technology, and — most importantly — what is required from EPROM to turn the "
    "prototype into a production platform."
)

# ---------------------------------------------------------------- 2. PROBLEM
eyebrow("Section 2")
h("The Problem With Today’s Approach")
body("Our current training relies almost entirely on three formats, each with real limitations:")
table(
    ["Current format", "Limitations"],
    [
        ["PowerPoint decks", "Passive; no interaction or feedback; quickly outdated; retention is low."],
        ["PDF manuals", "Dense and static; hard to visualise equipment; no way to assess understanding."],
        ["Physical site visits", "Expensive and slow; require travel, PPE and permits-to-work; safety exposure; limited seats; not repeatable on demand."],
    ],
    widths=[1.8, 4.5],
)
body("")
body(
    "The result is training that is costly to deliver, difficult to scale, inconsistent "
    "between groups, and provides no data on whether learning actually occurred."
)

# ---------------------------------------------------------------- 3. SOLUTION
eyebrow("Section 3")
h("The Solution: One Interactive Lesson")
body(
    "EPROM Academy replaces all three formats with a single interactive lesson player. "
    "The prototype demonstrates eight interaction types, any of which can be combined into a course:"
)
table(
    ["Interaction", "What it replaces / adds"],
    [
        ["Interactive content slides", "PowerPoint & PDF pages — self-paced, brand-styled, with diagrams."],
        ["Embedded video", "Process animations and field footage, streamed in context."],
        ["Interactive 3D equipment", "Cutaway diagrams — orbit, zoom and click hotspots to learn components."],
        ["360° virtual site visit", "Physical site visits — walk a plant area with no travel, PPE or permits."],
        ["Live equipment simulation", "Static curves — change operating conditions and watch performance respond."],
        ["Quizzes", "Verbal Q&A — instant feedback and automatic scoring."],
        ["Hands-on tasks", "On-the-job coaching — e.g. sequence a safe pump start-up."],
        ["Progress & completion", "Attendance sheets — per-learner tracking and certificates."],
    ],
    widths=[2.2, 4.1],
)
body("")
h("Demonstrated in the prototype", size=12, before=8)
body(
    "The sample course proves the model by replacing a 60-slide PowerPoint, a maintenance "
    "PDF and a half-day plant walk-down with one ~35-minute interactive lesson covering pump "
    "fundamentals, a 3D component tour, a virtual pump-house visit, a live performance "
    "simulator (including a cavitation warning), a knowledge-check quiz and a start-up "
    "sequencing task."
)

# ---------------------------------------------------------------- 4. ARCHITECTURE
eyebrow("Section 4")
h("Technical Architecture")
h("Prototype (built today)", size=12, before=8)
table(
    ["Layer", "Technology", "Notes"],
    [
        ["Frontend", "Next.js + React + TypeScript", "Modern, production-grade web framework."],
        ["Styling", "Tailwind CSS + EPROM brand tokens", "Matches the corporate visual identity."],
        ["3D & 360°", "three.js (React Three Fiber + drei)", "Runs in any modern browser, no plugins."],
        ["Simulation", "In-browser physics (JavaScript)", "Real-time pump performance model."],
        ["Data / progress", "Local mock data + browser storage", "No backend — keeps the demo simple."],
    ],
    widths=[1.4, 2.6, 2.3],
)
body("")
h("Production additions (to be built)", size=12, before=8)
bullet("User accounts & authentication (trainees, instructors, admins).")
bullet("Database to store courses, enrolments, progress, quiz results and certificates.")
bullet("Content management so subject-matter experts can author courses without code.")
bullet("Reporting dashboard (completion rates, scores, training records for audit/compliance).")
bullet("Media pipeline & hosting for video, 3D models and 360° imagery.")
bullet("Hosting, domain, backups and security hardening.")

# ---------------------------------------------------------------- 5. REQUIRED
eyebrow("Section 5")
h("What Is Required From EPROM")
body(
    "The technology is proven by the prototype. Turning it into a full platform depends "
    "mainly on content and a few organisational decisions:"
)
h("A. Content & people", size=12, before=8)
bullet("Subject-matter experts (SMEs) to provide course material and review accuracy.")
bullet("Source material: existing PowerPoint, PDFs and procedures to convert.")
bullet("Assessment questions and pass criteria for each course.")
h("B. Media assets", size=12, before=8)
bullet("Training videos (existing footage or newly produced process animations).")
bullet("3D / CAD models of key equipment (.glb/.gltf), or vendor models we convert.")
bullet("360° photography of EPROM sites — captured once per area, reused by everyone.")
h("C. Platform & operations", size=12, before=8)
bullet("Approval to build the backend (accounts, database, reporting).")
bullet("Hosting decision: EPROM servers, or a managed cloud (recommended for speed).")
bullet("A domain (e.g. academy.eprom.com.eg) and basic IT/security sign-off.")
h("D. Budget, team & timeline", size=12, before=8)
bullet("A small build team, or continued AI-assisted development, to deliver the MVP.")
bullet("Budget primarily driven by video production and 3D model creation (see Section 7).")
bullet("A pilot group of trainees to validate the first production course.")

# ---------------------------------------------------------------- 6. ROADMAP
eyebrow("Section 6")
h("Phased Roadmap")
table(
    ["Phase", "Scope", "Outcome"],
    [
        ["Phase 0 — Prototype", "Done. Sample course with all 8 interaction types, no backend.", "Proves the concept; this demo."],
        ["Phase 1 — MVP", "Accounts, database, one real EPROM course, progress tracking & certificate.", "First course live to a pilot group."],
        ["Phase 2 — Content scale-up", "Authoring tools; convert priority PPT/PDF courses; add real video.", "Library of interactive courses."],
        ["Phase 3 — Full immersion", "Real 3D/CAD models, 360° site captures, advanced simulations, admin reporting.", "Site visits & manuals fully replaced."],
    ],
    widths=[1.7, 3.3, 1.9],
)

# ---------------------------------------------------------------- 7. COST/RISK
eyebrow("Section 7")
h("Effort, Cost Drivers & Risks")
h("Main cost drivers", size=12, before=8)
bullet("Video production — the largest variable; can start with existing footage.")
bullet("3D / CAD model creation — reuse vendor models where possible to reduce cost.")
bullet("360° site photography — low cost; a one-time capture per area.")
bullet("Software build — contained, as the frontend is largely proven by the prototype.")
h("Key risks & mitigations", size=12, before=8)
bullet("SME availability — mitigate by converting existing material first.")
bullet("Asset creation time — phase it; launch with content slides + simulation, add 3D/360 later.")
bullet("Scope creep — the phased roadmap delivers value at each step before expanding.")

# ---------------------------------------------------------------- 8. DEMO GUIDE
eyebrow("Section 8")
h("How To Run & Present The Prototype")
numbered("Open a terminal in the project folder.")
numbered("Run “npm install” (first time only), then “npm run dev”.")
numbered("Open http://localhost:3000 in Chrome or Edge.")
numbered("Home → Browse catalog → open the Centrifugal Pump course → Start lesson.")
numbered("Step through with Next: content → video → 3D model (drag & click hotspots) → "
         "360° site visit → live simulation (move sliders; drop suction pressure to trigger "
         "cavitation) → quiz → start-up task → summary with score.")
body("")
note = doc.add_paragraph()
r = note.add_run(
    "Talking point for the demo: everything on screen — the 3D pump, the virtual site, "
    "the live curve — was built once and can be explored by every trainee, on any device, "
    "with no travel, PPE or permits, and with their results recorded automatically."
)
r.italic = True
r.font.color.rgb = MUTED

doc.save("EPROM-Academy-Proposal.docx")
print("Saved EPROM-Academy-Proposal.docx")
